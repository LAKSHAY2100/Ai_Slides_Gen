import base64
import csv
import io
import json
import os
from pathlib import Path
import random
import re
import string

from django.http import JsonResponse
from django.shortcuts import render
from django.views.decorators.csrf import csrf_exempt
from dotenv import load_dotenv
from google import genai
from google.genai import types
import tempfile
from urllib.parse import quote


from .models import SharedSlides

from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import OllamaEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import ChatPromptTemplate
from langchain_classic.chains.combine_documents import create_stuff_documents_chain
from langchain_community.llms import Ollama
from langchain_classic.chains import create_retrieval_chain
from langchain_groq import ChatGroq


load_dotenv()

os.environ["GROQ_API_KEY"] = os.getenv("GROQ_API_KEY")

GEMINI_API_KEY = os.getenv("GOOGLE_API_KEY")

client = genai.Client()
MAX_DOCUMENT_CHARS = 6000
ALLOWED_DOCUMENT_EXTENSIONS = {".pdf", ".csv", ".doc", ".docx"}


def slide_builder(request):
    return render(request, "slide_builder.html")


def _make_code(length=6):
    return "".join(random.choices(string.ascii_uppercase + string.digits, k=length))


@csrf_exempt
def share_slides(request):
    if request.method != "POST":
        return JsonResponse({"error": "Only POST Allowed"}, status=405)

    try:
        body = json.loads(request.body.decode("utf-8"))
        slides = body.get("slides", [])
    except Exception:
        return JsonResponse({"error": "Invalid Json"}, status=400)

    code = _make_code()
    SharedSlides.objects.create(share_code=code, slides_json=slides)
    return JsonResponse({"code": code})


def view_shared(request, code):
    try:
        record = SharedSlides.objects.get(share_code=code)
    except SharedSlides.DoesNotExist:
        return render(request, "not_found.html")

    return render(request, "slide_builder.html", {"shared_slides": json.dumps(record.slides_json)})


def _clean_source_text(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", text or "").strip()
    return cleaned[:MAX_DOCUMENT_CHARS]


def _extract_text_from_pdf(uploaded_file, topic="Random Topic", slide_type="General") -> str:
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
            for chunk in uploaded_file.chunks():
                tmp_file.write(chunk)
            tmp_path = tmp_file.name

        # 2. Load and Split
        loader = PyPDFLoader(tmp_path)
        docs = loader.load()
        
        # Cleanup the temp file immediately after loading into memory
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000 , chunk_overlap=20)
        documents =text_splitter.split_documents(docs)

        db = FAISS.from_documents(documents, OllamaEmbeddings())
        retriever = db.as_retriever(search_kwargs={"k": 8})

        llm=ChatGroq(model="openai/gpt-oss-120b")

        prompt = ChatPromptTemplate.from_template("""
        Summarize the context properly so that it can be used to generate slide titles and images about {topic} in a {slide_type} style presentation.
        <context>
        {context} 
        </context>"""
        )
        document_chain = create_stuff_documents_chain(llm , prompt)
        retrieval_chain = create_retrieval_chain(retriever, document_chain)

        response = retrieval_chain.invoke({"topic":topic,"slide_type":slide_type,"input": "Summarize this document"})

        print("Response from retrieval chain for PDF content extraction:", response['answer'])

        return response['answer']
    
    except Exception as e:
        print(f"Error in PDF extraction: {e}")
        # Fallback to basic extraction if RAG fails
        uploaded_file.seek(0)
        from pypdf import PdfReader
        reader = PdfReader(uploaded_file)
        return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_text_from_csv(uploaded_file) -> str:
    uploaded_file.seek(0)
    decoded = uploaded_file.read().decode("utf-8", errors="ignore")
    rows = []
    for row in csv.reader(io.StringIO(decoded)):
        cleaned_row = [cell.strip() for cell in row if cell.strip()]
        if cleaned_row:
            rows.append(" | ".join(cleaned_row))
        if len(rows) >= 40:
            break
    return "\n".join(rows)


def _extract_text_from_docx(uploaded_file) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("DOCX upload support is not installed on this server yet.") from exc

    uploaded_file.seek(0)
    document = Document(uploaded_file)
    sections = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                sections.append(" | ".join(cells))
    return "\n".join(sections)


def _extract_text_from_doc(uploaded_file) -> str:
    uploaded_file.seek(0)
    decoded = uploaded_file.read().decode("utf-8", errors="ignore")
    return re.sub(r"[^A-Za-z0-9\s,.;:!?()/@%&+\-]", " ", decoded)


def _extract_text_from_upload(uploaded_file,topic,slide_type) -> str:
    extension = Path(uploaded_file.name).suffix.lower()
    if extension not in ALLOWED_DOCUMENT_EXTENSIONS:
        allowed_types = ", ".join(sorted(ALLOWED_DOCUMENT_EXTENSIONS))
        raise ValueError(f"Unsupported file type. Please upload one of: {allowed_types}.")

    try:
        if extension == ".pdf":
            extracted = _extract_text_from_pdf(uploaded_file,topic,slide_type)
        elif extension == ".csv":
            extracted = _extract_text_from_csv(uploaded_file)
        elif extension == ".docx":
            extracted = _extract_text_from_docx(uploaded_file)
        else:
            extracted = _extract_text_from_doc(uploaded_file)
    except Exception as exc:
        raise ValueError(f"Could not read the uploaded {extension[1:].upper()} file.") from exc

    cleaned = _clean_source_text(extracted)
    if not cleaned:
        raise ValueError("The uploaded file did not contain readable text.")
    return cleaned


def _generate_slide_titles(topic: str, slide_type: str, source_text: str = "", document_name: str = "") -> list[str]:
    source_context = ""
    if source_text:
        source_context = f"""
        Use the uploaded document "{document_name or 'uploaded file'}" as the primary source.
        Prioritize its facts, headings, and structure when naming the slides.
        Source material:
        {source_text}
        """

    prompt = f"""
        Return exactly five slide titles for a beginner-friendly {slide_type} presentation about "{topic}".
        {source_context}
        Must return only valid JSON in this exact structure:
        {{
            "slides":[
                {{"title":"..."}},
                {{"title":"..."}},
                {{"title":"..."}},
                {{"title":"..."}},
                {{"title":"..."}}
            ]
        }}
    """

    try:
        response = client.models.generate_content(
            model="gemini-3-flash-preview",
            contents=[{"text": prompt}],
        )

        raw = ""
        for part in response.parts:
            if part.text:
                raw += part.text.strip()

        print("Raw title generation response:", raw)
        data = json.loads(raw)
        titles = [slide["title"] for slide in data.get("slides", []) if "title" in slide]

        if len(titles) == 5:
            return titles

    except Exception as exc:
        print("Title generation failed:", exc)

    return [
        f"Introduction to {topic}",
        f"Core ideas of {topic}",
        f"How {topic} works",
        f"Use case of {topic}",
        f"Future of {topic}",
    ]


def _generate_slide_image(title: str, topic: str, slide_type: str) -> str | None:

    prompt = (
        f"{slide_type} presentation slide about {topic}, focus on {title}, "
        "minimal, clean, soft colours, light background, no dense body text"
    )
    encoded_prompt = quote(prompt)
    print(encoded_prompt)
    POLLINATIONS_API_KEY = os.getenv("POLLINATIONS_API_KEY")
    return f"https://gen.pollinations.ai/image/{encoded_prompt}?model=flux&width=1280&height=720&nologo=true&key={POLLINATIONS_API_KEY}"

    prompt = (
        f"Create a slide of {slide_type} style with illustrations talking about {topic}."
        f"The focus of this slide is: {title}. "
        "Minimal, clean, soft colours on a light background, no dense body text."
    )

    print("Generating image with prompt :", prompt)

    try:
        response = client.models.generate_content(
            model="gemini-2.5-flash-image",
            contents=prompt,
            config=types.GenerateContentConfig(
                image_config=types.ImageConfig(
                    aspect_ratio="16:9",
                )
            ),
        )

        print("Image generation response : ", response)

        for part in response.parts:
            if part.inline_data:
                image_bytes = part.inline_data.data
                encoded = base64.b64encode(image_bytes).decode("ascii")
                mime_type = part.inline_data.mime_type or "image/jpeg"
                data_url = f"data:{mime_type}; base64,{encoded}"
                print("Generated image data url: ", data_url[:80], "...")
                return data_url

    except Exception as exc:
        print("Error in generate_slide_images :", exc)

    return None


@csrf_exempt
def generate_slides(request):
    if request.method != "POST":
        return JsonResponse({"error": "Only POST Allowed"}, status=405)

    topic = ""
    slide_type = ""
    document_text = ""
    document_name = ""
    content_type = request.content_type or ""

    if "application/json" in content_type:
        try:
            body = json.loads(request.body.decode("utf-8"))
        except json.JSONDecodeError:
            return JsonResponse({"error": "Invalid JSON"}, status=400)
        topic = body.get("topic", "").strip()
        slide_type = body.get("slide_type", "").strip()
    else:
        topic = request.POST.get("topic", "").strip()
        slide_type = request.POST.get("slide_type", "").strip()

    uploaded_document = request.FILES.get("document")
    if uploaded_document:
        document_name = uploaded_document.name
        try:
            document_text = _extract_text_from_upload(uploaded_document,topic,slide_type)
        except ValueError as exc:
            return JsonResponse({"error": str(exc)}, status=400)
        if not topic:
            topic = Path(document_name).stem.replace("_", " ").replace("-", " ").strip()

    if not topic:
        topic = "Random Topic"

    if not slide_type:
        slide_type = "General"

    titles = _generate_slide_titles(
        topic,
        slide_type,
        source_text=document_text,
        document_name=document_name,
    )
    print("Titles from AI:", titles)

    slides = []
    for idx, title in enumerate(titles):
        image_url = _generate_slide_image(title, topic, slide_type)
        if image_url is None:
            image_url = "https://images.unsplash.com/photo-1635070041078-e363dbe005cb?auto=format&fit=crop&w=800&q=80"

        slides.append(
            {
                "id": idx,
                "title": title,
                "image": image_url,
            }
        )

    return JsonResponse({"slides": slides})


def story_telling(request):
    return render(request, "story_telling.html")
